import json
import threading
from django.conf import settings
from django.shortcuts import render, redirect,get_object_or_404
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.mail import send_mail, EmailMultiAlternatives,EmailMessage
from django.http import JsonResponse,HttpResponse,Http404,FileResponse,HttpResponseNotAllowed
import os
import fitz
import docx
from django.utils.text import slugify
import mimetypes

# Core ScholarShield AI Module Imports
from .models import Profile,AcademicTask,AcademicNote
from .ask_ai import ask_ai
from .summarize_text_ai import ask_sum_text_ai
from .summarize_file import ask_sum_file_ai
from .task_planner import calculate_ai_priority
from .note_processor import process_note_with_ai
from .quiz_generator import generate_quiz_from_ai
# ==========================================
# AUTHENTICATION & LANDING VIEWS
# ==========================================

def index(request):
    if request.method == 'POST':
        first_name = request.POST.get("first_name", "").strip()
        last_name = request.POST.get("last_name", "").strip()
        email = request.POST.get("email", "").strip()
        department = request.POST.get("department", "").strip()
        level = request.POST.get("level", "").strip()
        password = request.POST.get("password", "").strip()

        # Validation logic
        if not first_name or not last_name or not email or not department or not level or not password:
            messages.error(request, "Field can't be empty.")
            print("❌ Registration Error: Fields cannot be empty.")
        else:
            # Check if user already exists
            if User.objects.filter(username=email).exists():
                messages.error(request, "User with this email already exists.")
                print("❌ Registration Error: User already exists.")
            else:
                user = User.objects.create_user(
                    username=email,
                    first_name=first_name,
                    last_name=last_name,
                    email=email,
                    password=password
                )
                Profile.objects.create(
                    user=user,
                    department=department,
                    level=level
                )
                print(f"✨ Success: User {first_name} created successfully!")
                
                # Non-blocking registration welcome email thread
                threading.Thread(
                    target=send_welcome_mail,
                    args=(email, first_name),
                    daemon=True
                ).start()
                
                messages.success(request, "User created successfully. Please log in.")
                return redirect("login")

    return render(request, "login/reg.html")


def login_view(request):
    if request.method == 'POST':
        email = request.POST.get("email", "").strip()
        password = request.POST.get("password", "")
        
        if not email or not password:
            messages.error(request, "Email or password can't be empty.")
            print("❌ Login Error: Missing credentials.")
        else:
            user = authenticate(request, username=email, password=password)
            
            if user is not None:
                login(request, user)
                messages.success(request, "Login successful")
                print(f"🔒 Session started for user: {user.email}")
                return redirect("dash")
            else:
                messages.error(request, "Invalid email or password.")
                print("❌ Login Error: Authentication failed.")

    return render(request, "login/login.html")


@login_required(login_url="login")
def Dash_board(request):
    all_active_tasks = AcademicTask.objects.filter(user=request.user, is_completed=False)
    existing_notes = AcademicNote.objects.filter(user=request.user).order_by('-id')
    critical_hotzone = all_active_tasks.filter(ai_score__gte=8.5)
    standard_feed = all_active_tasks.filter(ai_score__lt=8.5)
    context = {
        "critical_hotzone": critical_hotzone,
        "standard_feed": standard_feed,
        'notes': existing_notes,
    }
    return render(request, "login/dash.html",context)


def send_welcome_mail(user_email, name):
    subject = "Welcome to ScholarShield AI! 🎓"
    from_email = "ngadisuccess1@gmail.com"
    
    html_content = f"""
    <div style="font-family: Arial, sans-serif; padding: 20px; border: 1px solid #ccc;">
        <h1 style="color: #6a1b9a;">Welcome, {name}!</h1>
        <p>We are excited to have you join <strong>ScholarShield AI</strong>.</p>
        <p>Our platform is designed to make your academic life easier, safer, and much more productive.</p>
        <div style="margin-top: 20px; padding: 10px; background: #f9f9f9;">
            <p>Ready to start? Log in and explore your dashboard.</p>
        </div>
        <p>Best regards,<br>The Scientist Team, By success!😎</p>
    </div>
    """
    text_content = f"Hi {name}, welcome to ScholarShield AI! We are excited to have you."
    
    msg = EmailMultiAlternatives(subject, text_content, from_email, [user_email])
    msg.attach_alternative(html_content, "text/html")
    msg.send()


# ==========================================
# CORE CORE CORE - AI PIPELINES
# ==========================================

@login_required(login_url="login")
def chat_view(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            user_input = data.get("message")
            
            if 'chat_history' not in request.session:
                request.session['chat_history'] = []
                
            history = request.session['chat_history']
            
            ai_reply = ask_ai(
                user_input, 
                request.user.first_name, 
                history, 
                request.user.profile.department, 
                request.user.profile.level
            )
            
            history.append({"role": "user", "content": user_input})
            history.append({"role": "assistant", "content": ai_reply})
            
            request.session['chat_history'] = history
            request.session.modified = True 
            
            return JsonResponse({"reply": ai_reply})
        except Exception as e:
            print(f"❌ Chat View Crash: {str(e)}")
            return JsonResponse({"error": "Failed to process chat response."}, status=500)
            
    return JsonResponse({"error": "Only POST requests allowed"}, status=405)


@login_required(login_url="login")
def uploads_text(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            user_input = data.get("message")
            summary_format = data.get("summary_format", "bullet_points")
            custom_instructions = data.get("custom_instructions", "")

            ai_reply = ask_sum_text_ai(
                request.user.first_name, 
                user_input, 
                summary_format, 
                custom_instructions
            )
            return JsonResponse({"reply": ai_reply})
        except Exception as e:
            print(f"❌ Text Summarization Crash: {str(e)}")
            return JsonResponse({"error": "Failed to process text summary."}, status=500)
            
    return JsonResponse({"error": "Only POST requests allowed"}, status=405)


@login_required(login_url="login")
def uploads_file(request):
    if request.method == 'POST':
        uploaded_file = request.FILES.get('file')
        summary_format = request.POST.get('summary_format', 'bullet_points')
        custom_instructions = request.POST.get('custom_instructions', '')

        if not uploaded_file:
            return JsonResponse({'error': "No file stream received."}, status=400)

        file_name = uploaded_file.name
        
        try:
            ext = os.path.splitext(file_name)[1].lower()
            extracted_text = ""
            
            # Focused Lightweight Built-in Parsing for TXT files
            if ext == '.txt':
                print(f"📝 Processing plain-text upload: '{file_name}'")
                extracted_text = uploaded_file.read().decode('utf-8', errors='replace')
                
            elif ext == '.pdf':
                print(f"📄 Processing PDF via PyMuPDF: '{file_name}'")
                file_bytes = uploaded_file.read()
                with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                    text_chunks = []
                    for page in doc:
                        page_text = page.get_text()
                        if page_text:
                            text_chunks.append(page_text)
                            
                    extracted_text = "\n".join(text_chunks)
                
            elif ext in ['.docx', '.doc']:
                print(f"📘 Processing Word document via python-docx: '{file_name}'")
                doc = docx.Document(uploaded_file)
                text_chunks = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip(): 
                        text_chunks.append(paragraph.text)
                extracted_text = "\n".join(text_chunks)
            else:
                return JsonResponse({'error': "Unsupported file extension type."}, status=400)

            # Prevent empty file data payloads from breaking your Groq system
            if not extracted_text.strip():
                return JsonResponse({'error': "File content is completely empty."}, status=400)

            print(f"🚀 Text extraction successful. Forwarding payload to ask_sum_file_ai...")

            # Run LLM pipeline request execution
            ai_reply = ask_sum_file_ai(
                user_name=request.user.first_name or "Scientist",
                file_prompt=extracted_text,
                summary_format=summary_format,
                custom_instructions=custom_instructions
            )

            return JsonResponse({'reply': ai_reply})

        except Exception as e:
            # FIX: Captured and printed directly to terminal logs before returning response
            print("\n" + "="*60)
            print(f"🔥 SCHOLARSHIELD BACKEND PARSER CRASH: {str(e)}")
            print("="*60 + "\n")
            return JsonResponse({'error': f"Processing error: {str(e)}"}, status=500)

    return JsonResponse({'error': 'Invalid request method.'}, status=405)


@login_required(login_url="login")
def task_planner(request):
    # 1. Handle incoming AJAX/Fetch POST requests to add a task
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            title = data.get("title", "").strip()
            excerpt = data.get("excerpt", "").strip()
            tag = data.get("tag", "").strip()
            deadline_str = data.get("deadline", "").strip()

            if not title or not excerpt or not tag:
                return JsonResponse({"error": "Fields cannot be left empty."}, status=400)

            # Assign fallback deadline date if left blank
            if not deadline_str:
                from datetime import date, timedelta
                deadline_str = (date.today() + timedelta(days=3)).strftime("%Y-%m-%d")

            # 🔥 FIXED: Added real profile attributes to feed context parameters straight to Groq
            computed_score = calculate_ai_priority(
                title=title,
                excerpt=excerpt,
                deadline_date=deadline_str,
                user_name=request.user.first_name or "Scientist",
                department=request.user.profile.department,
                level=request.user.profile.level
            )

            # Save entity record directly inside the user's partition
            task = AcademicTask.objects.create(
                user=request.user,
                title=title,
                excerpt=excerpt,
                tag=tag,
                ai_score=computed_score,
                deadline=deadline_str
            )

            print(f"🎯 AI Priority Formulated: Task '{title}' assigned a rating of {computed_score}")
            return JsonResponse({
                "status": "success",
                "id": task.id,
                "title": task.title,
                "excerpt": task.excerpt,
                "tag": task.tag,
                "ai_score": task.ai_score,
                "deadline": str(task.deadline)
            })
        except Exception as e:
            print(f"❌ Planner Submission Failure: {str(e)}")
            return JsonResponse({"error": "Failed to store asset instance."}, status=500)

    # 2. Handle standard GET rendering requests
    all_active_tasks = AcademicTask.objects.filter(user=request.user, is_completed=False)
    
    # Pre-segment records directly inside the view to keep the HTML engine simple
    critical_hotzone = all_active_tasks.filter(ai_score__gte=8.5)
    standard_feed = all_active_tasks.filter(ai_score__lt=8.5)

    context = {
        "critical_hotzone": critical_hotzone,
        "standard_feed": standard_feed,
    }
    return render(request, "login/task_planner.html", context)


@login_required(login_url="login")
def resolve_task_endpoint(request, task_id):
    """
    Micro-service router used to flag active tasks as resolved asynchronously.
    """
    if request.method == "POST":
        try:
            task = AcademicTask.objects.get(id=task_id, user=request.user)
            task.is_completed = True
            task.save()
            return JsonResponse({"status": "resolved", "id": task_id})
        except AcademicTask.DoesNotExist:
            return JsonResponse({"error": "Task not found."}, status=404)
    return JsonResponse({"error": "Method not allowed."}, status=405)


@login_required(login_url="login")
def upload_note_file_view(request):
    if request.method == 'POST':
        uploaded_file = request.FILES.get('file')
        if not uploaded_file:
            return JsonResponse({'error': "No file uploaded."}, status=400)

        file_name = uploaded_file.name
        ext = os.path.splitext(file_name)[1].lower()
        extracted_text = ""
        
        # 1. Extraction Logic
        try:
            if ext == '.txt':
                extracted_text = uploaded_file.read().decode('utf-8', errors='replace')
            elif ext == '.pdf':
                with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
                    extracted_text = "\n".join([page.get_text() for page in doc])
            elif ext in ['.docx', '.doc']:
                doc = docx.Document(uploaded_file)
                extracted_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                return JsonResponse({'error': "Unsupported format."}, status=400)
        except Exception as e:
            return JsonResponse({'error': f"Extraction failed: {str(e)}"}, status=500)

        # 2. AI Processing
        if not extracted_text.strip():
            return JsonResponse({'error': "File content is empty."}, status=400)
            
        note_data = process_note_with_ai(extracted_text)
        print(note_data)
        
        # 3. Persistence
        new_note = AcademicNote.objects.create(
            user=request.user,
            topic=note_data['topic'],
            tags=note_data['tags'],
            ai_summary=note_data['summary'],
            raw_extracted_text=extracted_text,
            file_name=file_name,
        )
        
        return JsonResponse({
            "status": "success", 
            "id": new_note.id,
            "note": {
                "id": new_note.id,
                "topic": new_note.topic,
                "tags": new_note.tags,
                "ai_summary": new_note.ai_summary,  
                "file_name": new_note.file_name
            }
        })

@login_required(login_url="login")
def upload_note_text_view(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            raw_text = data.get("message", "").strip()
            
            if not raw_text:
                return JsonResponse({"error": "No text provided."}, status=400)
            
            note_data = process_note_with_ai(raw_text)
            
            new_note = AcademicNote.objects.create(
                user=request.user,
                topic=note_data['topic'],
                tags=note_data['tags'],
                ai_summary=note_data['summary'],
                raw_extracted_text=raw_text,
                file_name="Pasted Note"
            )
            
            return JsonResponse({
                "status": "success", 
                "id": new_note.id,
                "note": {
                    "id": new_note.id,
                    "topic": new_note.topic,
                    "tags": new_note.tags,
                    "ai_summary": new_note.ai_summary,  # This MUST match the JS expectation
                    "file_name": new_note.file_name
                }
            })
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

@login_required(login_url="login")
def delete_note_view(request, note_id):
    if request.method == "POST":
        try:
            AcademicNote.objects.get(id=note_id, user=request.user).delete()
            return JsonResponse({"status": "deleted"})
        except AcademicNote.DoesNotExist:
            return JsonResponse({"error": "Not found."}, status=404)



@login_required
def download_note(request, note_id):
    note = get_object_or_404(AcademicNote, id=note_id, user=request.user)


    if note.uploaded_file:
        filename = os.path.basename(note.uploaded_file.name)
        resp = FileResponse(note.uploaded_file.open("rb"), as_attachment=True, filename=f"{filename}-Scholar-shieldAI{note_id}")
        return resp

    content = note.raw_extracted_text or note.ai_summary or ""
    if not content.strip():
        raise Http404("Nothing to download for this note.")

    safe_topic = slugify(note.topic or "note")[:50] or "note"
    filename = f"{safe_topic}.txt"  

    response = HttpResponse(content, content_type="text/plain; charset=utf-8")
    response["Content-Disposition"] = f'attachment; filename=f"{filename}-Scholar-shieldAI{note_id}"'
    return response



def _send_note_to_self_email(note_id, user):
    note = AcademicNote.objects.get(id=note_id, user=user)

    recipient = (user.email or "").strip()
    if not recipient:
        return  # no email to send to

    from_email = getattr(settings, "DEFAULT_FROM_EMAIL", None)

    subject = f"ScholarShield AI Note: {note.topic}"
    body = (
        f"Hi {user.first_name or user.username},\n\n"
        f"Thank you for user Scholar shield AI. \n\n"
        f"Here is your note.\n"
        f"Topic: {note.topic}\n"
        f"Label: {note.file_name}\n\n"
        f"— ScholarShield AI"
    )

    msg = EmailMessage(subject, body, from_email, [recipient])

    if note.uploaded_file:
        filename = os.path.basename(note.uploaded_file.name)
        mime = mimetypes.guess_type(filename)[0] or "application/octet-stream"
        with note.uploaded_file.open("rb") as f:
            msg.attach(filename, f.read(), mime)
    else:
        content = note.raw_extracted_text or note.ai_summary or ""
        if not content.strip():
            return
        safe_topic = slugify(note.topic or "note")[:50] or "note"
        filename = f"{safe_topic}-ScholarShieldAI-{note.id}.txt"
        msg.attach(filename, content.encode("utf-8"), "text/plain")

    msg.send(fail_silently=False)


@login_required
def share_note(request, note_id):
    note = get_object_or_404(AcademicNote, id=note_id, user=request.user)

    if not (request.user.email or "").strip():
        return JsonResponse({"error": "No email address found on your account."}, status=400)

    # send in background so the request returns quickly (no page reload)
    threading.Thread(
        target=_send_note_to_self_email,
        args=(note.id, request.user),
        daemon=True
    ).start()

    return JsonResponse({"status": "sent", "to": request.user.email, "note_id": note.id})


@login_required
def get_note_data(request, note_id):
    note = get_object_or_404(AcademicNote, id=note_id, user=request.user)

    return JsonResponse({
        "id": note.id,
        "file_name": note.file_name,
        "has_file": bool(note.uploaded_file),
        "file_url": note.uploaded_file.url if note.uploaded_file else None,
        "raw_text": note.raw_extracted_text or "",
    })



@login_required
def generate_quiz(request, note_id):
    note = get_object_or_404(
        AcademicNote, 
        id=note_id, 
        user=request.user
    )

    raw_text = note.raw_extracted_text or ""
    if not raw_text.strip():
        return JsonResponse(
            {"error": "Note has no content."}, 
            status=400
        )

    quiz_data = generate_quiz_from_ai(raw_text)

    return JsonResponse({"quiz": quiz_data})